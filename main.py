# coding=utf-8
import pdfplumber
import requests
import os
import docx
from functools import reduce
import re
import datetime
import pandas as pd
import platform
import subprocess
from logging import handlers
import logging
import json

"""""""""""""""""""""""""""

接口功能函数

"""""""""""""""""""""""""""


# 日志函数，直接添加使用即可
def create_logger(log_path):
    level_relations = {
        'debug': logging.DEBUG,
        'info': logging.INFO,
        'warning': logging.WARNING,
        'error': logging.ERROR,
        'crit': logging.CRITICAL
    }  # 日志级别关系映射

    logger = logging.getLogger(log_path)
    fmt = '%(asctime)s - %(pathname)s[line:%(lineno)d] - %(levelname)s: %(message)s'
    format_str = logging.Formatter(fmt)  # 设置日志格式
    logger.setLevel(level_relations.get('info'))  # 设置日志级别
    sh = logging.StreamHandler()  # 往屏幕上输出
    sh.setFormatter(format_str)  # 设置屏幕上显示的格式
    th = handlers.TimedRotatingFileHandler(
        filename=log_path, when='D', backupCount=3,
        encoding='utf-8')  # 往文件里写入#指定间隔时间自动生成文件的处理器
    th.setFormatter(format_str)  # 设置文件里写入的格式
    logger.addHandler(sh)  # 把对象加到logger里
    logger.addHandler(th)
    return logger


# 开始记录日志
# root_dir = os.path.dirname(__file__)
root_dir = os.getcwd()  # 获取当前根目录
logger = create_logger(root_dir + '/log/' + datetime.datetime.now().strftime('%Y-%m-%d') + '.log')  # 日志文件位置

def extract_interface(text):
    """
    function: 调用抽取模型接口
    param: 需要模型抽取的一段文字
    return: 抽取结果，json格式
    """
    try:
        url = "https://playground.datacanvas.com/aps/msserver/project/96aa5476-5d9b-4d12-8c83-fe6844897f9c/service/df8bf5ea-525f-45a4-a2f9-b0a87683d653/predict"
        headers = {'Content-Type': 'application/json;format=v1', }
        p = {"apiKey": "52dc0e3d-dcde-4df3-ae36-3a814a253762",
             "requestId": "e04714a1-2c03-4253-9a98-d1888d6a31c3",
             "data": [[
                 {"name": "text",
                 "type": "string",
                 "value": text,
                 }]]}
        p = json.dumps(p)
        res = requests.post(url, headers=headers, data=p)
        result = res.json()['data'][0]['value']
    except:
        result= {'无抽取实体': '无'}
    return result


def extract_func(text):
    """
    function: 模型抽取功能
    param: 文本
    return: 抽取结果
    """
    # 将文本按句号切分，对每一个句子过一遍模型，结果整合在result_all中
    sp_datas = text.split('。')
    result_all = []
    for sp_data in sp_datas:
        if len(sp_data) == 0:
            continue
        else:
            if len(sp_data) < 500:  # 单次抽取文本长度不能太长，太长则跳过
                result = extract_interface(sp_data)
                print('result',result)
                result_all.append(result)
            else:
                logger.info('------存在单句结果过长：{}'.format(sp_data))
    # 对抽取结果result_all去除value为None的key，结果保存在dup_res中
    dup_res=[]
    for result_sin in result_all:
        result_sin=eval(result_sin)
        if len(result_sin) != 0:
            dup_res.append(result_sin)
    # 将dup_res中的大字典拆分，使每个小字典只有一个键值对，以保证后面最大程度得去除重复元素，结果保存在splitdic中
    splitdic = []
    for i in dup_res:
        if len(i) == 1:
            splitdic.append(i)
        else:
            for key, value in i.items():
                small = {}
                small[key] = value
                splitdic.append(small)
    logger.info('步骤一：{}'.format(splitdic))
    # 对抽取结果去除键值对完全一样的项，结果保存在res中
    run_function = lambda x, y: x if y in x else x + [y]
    dup_res = reduce(run_function, [[], ] + splitdic)
    res = {}
    for i in dup_res:
        for key, value in i.items():
            if key not in res:
                res[key] = value
            else:
                res[key] = res[key] + '，' + value
    logger.info('步骤二：{}'.format(res))
    for key, value in res.items():
        mid_value = value.split('，')
        run_function = lambda x, y: x if y in x else x + [y]
        dup_res = reduce(run_function, [[], ] + mid_value)
        res[key] = '，'.join(dup_res)
    logger.info('步骤三：{}'.format(res))

    return res


# ocr接口
def ocr_interface(pdf_name, path):
    """
    function: 调用ocr接口
    param: pdf文件名称，pdf文件路径
    return:res为ocr抽取的文字结果
    """
    res = []
    url = "http://116.31.82.218:18088/nlpApi/ocrPtTransaction/ocrApi/ocrNlp"
    payload = {'sysId': 'NLP',
               'accesskey': '831b76ba311d40eda29d398d5cb4a3a8',
               'billType': 'FULLTEXT'}
    files = [('imgs', (pdf_name, open(path, 'rb'), 'application/pdf'))]
    headers = {}
    response = requests.request("POST", url, headers=headers, data=payload, files=files)
    # 获取pdf中的文字，将文字连接起来，替换掉\n和\r
    if response.json():
        for i in response.json()['icrResults']:
            try:
                res.append(i['results'].replace('\n', '').replace('\r', ' '))
            except:
                pass
    return res


# 判断pdf是否为图片
def is_img_page(page):
    """
    function: 判断传入的page对象是否为图片型page
    param: pdf中的每一页
    return: 是否为图片型
    """
    chars = page.chars
    # 无文字则为图片型
    if len(chars) == 0:
        return True
    elif len(page.images) == 0:
        # 无图片则为非图片型。
        return False
    elif len(page.images) > 0:
        page_area = page.width * page.height
        img_total_area = 0
        for img in page.images:
            img_total_area += img['width'] * img['height']
        # 如果图片累计面积超过页面的1/10
        if img_total_area > 0.1 * page_area:
            return True
        else:
            return False


# doc转docx
def doc_docx(this_path, doc_path, file_name):
    """
    function: doc文件在抽取之前需要先把doc格式转换为docx
    param: doc文件的根目录路径、doc转换之后的存储路径、文件名称
    return: 无输出，直接转存为docx文件
    """
    save_path = doc_path + file_name
    # 系统是linux，使用libreoffice
    if platform.system() == "Linux":
        try:
            args = ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', doc_path, this_path]
            process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=None)
            re.search('-> (.*?) using filter', process.stdout.decode())
            logger.info('---------------------------doc文件linux环境下已经转为docx：{}'.format(doc_path + file_name + 'x'))
        except:
            logger.info('---------------------------doc文件linux环境下转存报错：{}'.format(this_path))
    # 系统是win，使用win32com
    elif platform.system() == "Windows":
        try:
            import win32com.client
            word = win32com.client.Dispatch('kwps.Application')
            doc = word.Documents.Open(this_path)  # 目标路径下的文件
            doc.SaveAs(save_path + 'x', 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
            doc.Close()
            word.Quit()
            logger.info('---------------------------doc文件wins环境下已经转为docx：{}'.format(doc_path + file_name + 'x'))
        except:
            logger.info('---------------------------doc文件wins环境下转存报错：{}'.format(this_path))


def read_ocr_pdf(ocr_text, filename):
    """
    function: ocrpdf文书读取-处理-提取-输出
    param: ocr抽取的文字，文件名
    return: table_r为ocr文字的表格抽取结果（为空）、text_res为文字抽取结果
    """
    table_res, text_res = [], []  # 用于存储结果
    logger.info('-----------图片型pdf读取完毕，开始调用接口抽取信息{}'.format(filename))
    texts = ''.join(ocr_text)  # 用句号分割文字
    texts = texts.replace(' ', '')

    logger.info('-----------开始正则抽取：{}'.format(filename))
    regular_res = regular_extraction(texts, filename)  # 正则抽取结果
    logger.info('------正则抽取结果：{}'.format(regular_res))

    logger.info('----------开始模型抽取：{}'.format(filename))
    try:
        model_res = extract_func(texts)  # 模型抽取结果
    except:
        model_res= {'无抽取实体': '无'}
    logger.info('------模型抽取结果：{}'.format(model_res))

    # 如果模型抽取不出东西。把{'无抽取实体':'无'}作为模型的抽取结果
    if len(model_res) == 0:
        model_res = {'无抽取实体': '无'}
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    # 如果模型可以抽取出东西。把模型的抽取结果和正则抽取的结果做拼接
    else:
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    for key, value in text_res[0].items():
        if value == '' or value == None:
            text_res[0][key] = '无'
    logger.info('-------------抽取结果：{}'.format(text_res))

    return table_res, text_res


def read_pdf(text_path, filename):
    """
    function: pdf文书读取-处理-提取-输出
    param: pdf路径，文件名
    return: table_res为pdf的表格抽取结果、text_res为文字抽取结果
    """
    logger.info('-----------文字型pdf开始读取：{}'.format(filename))
    read_text = []  # 存储pdf读取的所有文字（不包括表格文字）
    clist = []  # 暂存，待删
    table_res, text_res = [], []  # 用于存储结果
    pdf = pdfplumber.open(text_path)
    for i in range(len(pdf.pages)):  # 读取每一页
        if pdf.pages[i].extract_tables():  # 如果存在表格，先将表格抽取到csv文件
            clist.append(pdftable_csv(pdf.pages[i]))

        for j in pdf.pages[i].extract_words():  # 读取每页中的每一行的文字，放入列表中
            if j['text']:
                read_text.append(j['text'])
    table_res.append(clist)
    logger.info('-----------读取完毕，开始调用接口抽取信息：{}'.format(filename))
    read_text = read_text + [_ for _ in flat(table_res)]  # 将表格文字与其他部分合并，然后一起抽取
    texts = ''.join(read_text)  # 用句号分割文字
    texts = texts.replace(' ', '')

    logger.info('-----------文字型pdf开始信息抽取：{}'.format(filename))
    logger.info('-----------开始正则抽取：{}'.format(filename))
    regular_res = regular_extraction(texts, filename)  # 正则抽取结果
    logger.info('------正则抽取结果：{}'.format(regular_res))

    logger.info('-----------开始模型抽取：{}'.format(filename))
    try:
        model_res = extract_func(texts)  # 模型抽取结果
    except:
        model_res= {'无抽取实体': '无'}
    logger.info('------模型抽取结果：{}'.format(model_res))
    # 如果模型抽取不出东西。把{'无抽取实体':'无'}作为模型的抽取结果
    if len(model_res) == 0:
        model_res = {'无抽取实体': '无'}
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    # 如果模型可以抽取出东西。把模型的抽取结果和正则抽取的结果做拼接
    else:
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    for key, value in text_res[0].items():
        if value == '' or value == None:
            text_res[0][key] = '无'
    logger.info('-------------抽取结果：{}'.format(text_res))
    return table_res, text_res


def read_word(text_path, filename):
    """
    function: word文书读取-处理-提取-输出
    param: word路径，文件名
    return: table_res为word文字的表格抽取结果、text_res为文字抽取结果
    """
    logger.info('------------docx文件开始读取：{}'.format(filename))
    read_text = []  # 存储pdf读取的所有文字（不包括表格文字）
    table_res, text_res = [], []  # 用于存储结果
    # 获取文档对象
    file = docx.Document(text_path)
    if file.tables:  # 如果存在表格，先抽取表格
        table_res.append(wordtable_csv(file))
    # 输出每一段的内容，存入列表
    for para in file.paragraphs:
        if para.text:
            read_text.append(para.text)

    logger.info('------------读取完毕，开始调用接口抽取信息：{}'.format(filename))
    read_text = read_text + [_ for _ in flat(table_res)]
    texts = ''.join(read_text)
    texts = texts.replace('\t', '')

    logger.info('------------开始正则抽取：{}'.format(filename))
    regular_res = regular_extraction(texts, filename)  # 正则抽取结果
    logger.info('------正则抽取结果：{}'.format(regular_res))

    logger.info('------------开始模型抽取：{}'.format(filename))
    try:
        model_res = extract_func(texts)  # 模型抽取结果
    except:
        model_res= {'无抽取实体': '无'}
    logger.info('------模型抽取结果：{}'.format(model_res))
    # 如果模型抽取不出东西。把{'无抽取实体':'无'}作为模型的抽取结果
    if len(model_res) == 0:
        model_res = {'无抽取实体': '无'}
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    # 如果模型可以抽取出东西。把模型的抽取结果和正则抽取的结果做拼接
    else:
        text_res = result_concat(model_res, regular_res)  # 结果拼接
    for key, value in text_res[0].items():
        if value == '' or value == None:
            text_res[0][key] = '无'
    logger.info('-------------抽取结果：{}'.format(text_res))
    return table_res, text_res


# 读取word表格
def wordtable_csv(file):
    """
    function: 读取word中的表格
    param: docx文件名
    return: 获取的表格结果，嵌套列表
    """
    result = []
    # 读取docx文件中的表格
    for table in file.tables:
        clist = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]  # 解析后的内容往这个空列表里面存
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    clist[i][j] = cell.text
        result.append(clist)
    return result


# 读取pdf表格
def pdftable_csv(pages):
    """
    function: 读取pdf中的表格
    param: pdf文件名
    return: 获取的表格结果，嵌套列表
    """
    result = []
    tables = pages.extract_tables()
    for table in tables:  # 遍历表格
        for row in table:  # 遍历表格中的每一行
            row = [x.strip() for x in row if x != None]
            result.append(row)
    return result


# 展平
def flat(l):
    """
    function: 拉平嵌套函数，用于获取表格中的内容
    param: 需要拉平的列表
    return: 拉平后的列表
    """
    for _ in l:
        if isinstance(_, list):
            yield from flat(_)
        else:
            yield _


def result_concat(model_result, regular_result):
    """
    function: 将模型抽取结果和正则表达式抽取结果进行整合
    param: model_result为模型抽结果,regular_result为正则表达式抽取结果
    return: reg_res1为在正则结果基础上经过模型结果补充的最终抽取结果
    """
    # 如果正则抽取结果regular_result为空，说明没能判断该文书是哪种类型，则将模型抽取结果model_result直接赋给变量reg_res作为返回值输出
    if len(regular_result) == 0:
        reg_res = [model_result]
        logger.info('-------------合并结果（正则抽取为空）：{}'.format(reg_res))
    # 如果正则抽取结果regular_result非空则按照如下逻辑处理
    else:
        # 将模型结果、正则结果分别放入model_res3和reg_res中
        mid_res = []
        for key, value in model_result.items():
            mid_res.append({key: value})
        for key, value in regular_result.items():
            if value == None:
                regular_result[key] = ''
        reg_res = [regular_result]
        model_res = tuple(mid_res)
        # 对于model_res3中每一个模型抽取结果mid，mid为字典，如{'被告': '张永红'}
        for i in model_res:
            mid = i
            for j in list(mid.keys()):
                # 如果key例如‘被告’是正则结果reg_res[0]中存在（即需要抽取）的key且reg_res[0]抽取的value非空，则从mid中删除该键值对
                if j in reg_res[0].keys() and len(reg_res[0].get(j)) != 0:
                    del i[j]
                # 如果key例如‘上诉人’不是正则结果reg_res[0]中存在的key（说明这一项不是该文书要抽取的），则从mid中删除该键值对
                elif j not in reg_res[0].keys():
                    del i[j]
                else:
                    continue
        # 把model_res3中所剩的项放入lost中，然后更新到正则结果reg_res[0]中，并作为返回值输出
        lost = []
        for i in model_res:
            if i != {}:
                lost.append(i)
        for i in range(0, len(lost)):
            reg_res[0].update(lost[i])
        logger.info('-------------合并结果：{}'.format(reg_res))
    return reg_res


# docx结果清洗保存
def result_process_docx(path, filename):
    """
    function: docx文件最终的结果清洗
    param: 文字抽取结果、表格抽取结果、抽取目标列表、结果保存路径、文件名
    return: 存储成csv，返回文字抽取结果、表格抽取结果
    """
    dict_result, end_result = [], []
    table_result, text_exresult = read_word(path, filename)
    # 对text_exresult[0]中的键值对，如果存在某一项的值例如 ‘被告’：【‘张三’，‘李四’】，则把value合并到一个字符串中，用，分隔例如‘被告’：‘张三，李四’
    try:
        for k, v in text_exresult[0].items():
            if type(v) == list:
                v1 = '，'.join(v)
                text_exresult[0][k] = v1
        end_result = [text_exresult[0]]
    except:
        logger.info('---------------------------抽取结果清洗失败')
    # 写入抽取的表格
    try:
        if len(table_result) == 0:
            table_result = ['无']
        else:
            for i in table_result[0]:  # 可能存在多个表格
                if len(i) != 0:
                    dict_result.append(pd.DataFrame(i[1:], columns=i[0]).to_dict('list'))
    except:
        logger.info('---------------------------表格结果清洗失败')
    return dict_result, end_result


# pdf结果清洗保存
def result_process_pdf(path, filename):
    """
    function: pdf文件最终的结果清洗
    param: 文字抽取结果、表格抽取结果、抽取目标列表、结果保存路径、文件名
    return: 存储成csv，返回文字抽取结果、表格抽取结果
    """
    logger.info('----------文字型pdf开始抽取：{}'.format(path))
    dict_result, end_result = [], []
    table_result, text_exresult = read_pdf(path, filename)
    # 对text_exresult[0]中的键值对，如果存在某一项的值例如 ‘被告’：【‘张三’，‘李四’】，则把value合并到一个字符串中，用，分隔例如‘被告’：‘张三，李四’
    try:
        for k, v in text_exresult[0].items():
            if type(v) == list:
                v1 = '，'.join(v)
                text_exresult[0][k] = v1
        end_result = [text_exresult[0]]
    except:
        logger.info('---------------------------抽取结果清洗失败')
    try:
        # 写入抽取的表格
        if len(table_result) == 0:
            table_result = ['无']
        else:
            for i in table_result[0]:  # 可能存在多个表格
                if len(i) != 0:
                    dict_result.append(pd.DataFrame(i[1:], columns=i[0]).to_dict('list'))
    except:
        logger.info('---------------------------表格结果清洗失败')
    return dict_result, end_result


"""""""""""""""""""""""""""

正则表达式抽取

"""""""""""""""""""""""""""


# 判断文书种类
def regular_extraction(text_str, filename):
    """
    function: 通过正则判断该文书是什么类型
    param: 读入的文书文本
    return: 不同的文书类型调用不同的信息抽取函数，如ex_civil_judgement( )为民事判决书抽取函数，返回抽取结果clas
    """
    text_str = text_str.replace(' ', '')
    text_str = text_str.replace('\n', '')
    try:
        # 抽取文本的前30个字，通过a~g的值判别文书对应的类型
        str2 = text_str[:30]
        str2 = re.sub(r"([^\u4e00-\u9fa5])", "", str2)
        a = re.search(r'[\u4e00-\u9fa5]+人民法院民事判决书', str2)
        b = re.search(r'民事起诉状', str2)
        c = re.search(r'民事上诉状', str2)
        d = re.search(r'[\u4e00-\u9fa5]+人民法院举证通知书', str2)
        e = re.search(r'[\u4e00-\u9fa5]+人民法院应诉通知书', str2)
        f = re.search(r'[\u4e00-\u9fa5]+异议起诉状', str2)
        g = re.search(r'委托代理合同', str2)
        tag1 = (0, 1)[a != None]
        tag2 = (0, 1)[b != None]
        tag3 = (0, 1)[c != None]
        tag4 = (0, 1)[d != None]
        tag5 = (0, 1)[e != None]
        tag6 = (0, 1)[f != None]
        tag7 = (0, 1)[g != None]
        tagdic = {1: tag1, 2: tag2, 3: tag3, 4: tag4, 5: tag5, 6: tag6, 7: tag7}
        clsdic = {1: '此文本为民事判决书', 2: '此文本为民事起诉状', 3: '此文本为民事上诉状', 4: '此文本为举证通知书', 5: '此文本为应诉通知书', 6: '此文本为执行异议起诉状',
                  7: '此文本为委托代理合同'}
        if sum(tagdic.values()) == 0:
            # 如果clsdic中a~g的值都为0，说明靠文本标题无法判断该文书类型，则使用文件名在进行一次判断
            try:
                aa = re.search(r'判决', filename)
                bb = re.search(r'案起诉状', filename)
                cc = re.search(r'上诉状', filename)
                dd = re.search(r'举证通知书', filename)
                ee = re.search(r'应诉通知书', filename)
                ff = re.search(r'异议起诉状', filename)
                gg = re.search(r'委托代理合同', filename)
                tag11 = (0, 1)[aa != None]
                tag22 = (0, 1)[bb != None]
                tag33 = (0, 1)[cc != None]
                tag44 = (0, 1)[dd != None]
                tag55 = (0, 1)[ee != None]
                tag66 = (0, 1)[ff != None]
                tag77 = (0, 1)[gg != None]
                tagdic2 = {1: tag11, 2: tag22, 3: tag33, 4: tag44, 5: tag55, 6: tag66, 7: tag77}
                clsdic2 = {1: '此文本为民事判决书', 2: '此文本为民事起诉状', 3: '此文本为民事上诉状', 4: '此文本为举证通知书', 5: '此文本为应诉通知书',
                           6: '此文本为执行异议起诉状', 7: '此文本为委托代理合同'}
                if sum(tagdic2.values()) == 0:
                    # 说明文书标题和文件名均无法判断
                    logger.info('-------------不属于任意类型：{}'.format(filename))
                    clas = {}
                    return clas
                elif sum(tagdic2.values()) > 1:
                    # 判断出文书属于不止一种类型
                    logger.info('-------------文书属于不止一种类型，判断出错：{}'.format(filename))
                    clas = {}
                    return clas
                else:
                    # 文件名判断出了文件类型，调用对应文书类型的正则抽取函数
                    index = [k for k, v in tagdic2.items() if v == 1][0]
                    cls = clsdic2.get(index)
                    if cls == '此文本为民事判决书':
                        clas = ex_civil_judgement(text_str)
                        return clas
                    elif cls == '此文本为民事起诉状':
                        clas = ex_civil_complaint(text_str)
                        return clas
                    elif cls == '此文本为民事上诉状':
                        clas = ex_civil_appeal_petition(text_str)
                        return clas
                    elif cls == '此文本为举证通知书':
                        clas = ex_proof(text_str)
                        return clas
                    elif cls == '此文本为应诉通知书':
                        clas = ex_respondence(text_str)
                        return clas
                    elif cls == '此文本为执行异议起诉状':
                        clas = ex_obj2exe(text_str)
                        return clas
                    else:
                        clas = ex_contract(text_str)
                        return clas
            except:
                pass
        elif sum(tagdic.values()) > 1:
            # 判断出文书属于不止一种类型
            logger.info('-------------文书属于不止一种类型，判断出错：{}'.format(filename))
            clas = {}
            return clas
        else:
            # 文本标题判断出了文件类型，调用对应文书类型的正则抽取函数
            index = [k for k, v in tagdic.items() if v == 1][0]
            cls = clsdic.get(index)
            if cls == '此文本为民事判决书':
                clas = ex_civil_judgement(text_str)
                return clas
            elif cls == '此文本为民事起诉状':
                clas = ex_civil_complaint(text_str)
                return clas
            elif cls == '此文本为民事上诉状':
                clas = ex_civil_appeal_petition(text_str)
                return clas
            elif cls == '此文本为举证通知书':
                clas = ex_proof(text_str)
                return clas
            elif cls == '此文本为应诉通知书':
                clas = ex_respondence(text_str)
                return clas
            elif cls == '此文本为执行异议起诉状':
                clas = ex_obj2exe(text_str)
                return clas
            else:
                clas = ex_contract(text_str)
                return clas
    except:
        logger.info('---------------------------判断文书种类失败')


# 民事判决书调用函数
def ex_civil_judgement(text_str):
    """
    function: 判断为民事判决书的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值civil_judgement为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_defendant(text_str)
    entity_2 = find_plaintiff(text_str)
    entity_3 = find_casecode(text_str)
    entity_4 = find_claim(text_str)
    entity_5 = find_judgment(text_str)
    entity_6 = find_court1(text_str)
    civil_judgement = {'被告': entity_1, '原告': entity_2, '案号': entity_3, '诉讼请求': entity_4, '判决结果': entity_5,
                       '法院名称': entity_6}
    return civil_judgement


# 民事起诉状调用函数
def ex_civil_complaint(text_str):
    """
    function: 判断为民事起诉状的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值civil_complaint为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_defendant(text_str)
    entity_2 = find_plaintiff(text_str)
    entity_3 = find_claim(text_str)
    entity_4 = find_court2(text_str)
    civil_complaint = {'被告': entity_1, '原告': entity_2, '诉讼请求': entity_3, '法院名称': entity_4}
    return civil_complaint


# 民事上诉状调用函数
def ex_civil_appeal_petition(text_str):
    """
    function: 判断为民事上诉状的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值civil_petition为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_appellee(text_str)
    entity_2 = find_appellor(text_str)
    entity_3 = find_claim(text_str)
    entity_4 = find_court2(text_str)
    entity_5 = find_date()
    civil_petition = {'被上诉人': entity_1, '上诉人': entity_2, '上诉请求': entity_3, '法院名称': entity_4, '日期': entity_5}
    return civil_petition


# 举证通知书
def ex_proof(text_str):
    """
    function: 判断为举证通知书的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值proof为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_casename1(text_str)
    entity_2 = find_casecode(text_str)
    entity_3 = find_bankname(text_str)
    entity_4 = find_court1(text_str)
    entity_5 = find_date()
    proof = {'案件名称': entity_1, '案号': entity_2, '抬头银行名称': entity_3, '法院名称': entity_4, '时间': entity_5}
    return proof


# 应诉通知书
def ex_respondence(text_str):
    """
    function: 判断为应诉通知书的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值respondence为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_casename1(text_str)
    entity_2 = find_casecode(text_str)
    entity_3 = find_bankname(text_str)
    entity_4 = find_court1(text_str)
    entity_5 = find_date()
    respondence = {'案件名称': entity_1, '案号': entity_2, '抬头银行名称': entity_3, '法院名称': entity_4, '时间': entity_5}
    return respondence


# 执行异议起诉状
def ex_obj2exe(text_str):
    """
    function: 判断为执行异议起诉状的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值obj2exe为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_defendant(text_str)
    entity_2 = find_plaintiff(text_str)
    entity_3 = find_third(text_str)
    entity_4 = find_claim(text_str)
    entity_5 = find_truth(text_str)
    entity_6 = find_date()
    obj2exe = {'被告': entity_1, '原告': entity_2, '第三人': entity_3, '诉讼请求': entity_4, '事实与理由': entity_5, '日期': entity_6}
    return obj2exe


# 委托代理合同
def ex_contract(text_str):
    """
    function: 判断为委托代理合同的文本信息抽取函数
    param: 读入的文书文本
    return: 返回值ex_contract为该文本抽取结果的{关键信息实体：实体内容}键值对
    """
    entity_1 = find_partyA(text_str)
    entity_2 = find_partyB(text_str)
    entity_case = ''.join(entity_1)
    entity_3 = find_casename2(text_str)
    entity_4 = find_lawyers(text_str)
    entity_5 = find_cost(text_str)
    entity_6 = find_date()
    ex_contract = {'甲方': entity_1, '乙方': entity_2, '案件名称': ''.join((entity_case, entity_3)), '代理律师': entity_4,
                   '费用': entity_5, '日期': entity_6}
    return ex_contract


"""""""""""""""""""""""""""

具体关键信息抽取函数

"""""""""""""""""""""""""""
# 被告
def find_defendant(text_str):
    """
    function: 被告信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的被告信息
    """
    text_str = text_str.replace(' ', '')
    try:
        plaintiff = re.findall(r'被告[0-9]{0,1}[：:].*?[,，住所地]+', text_str)
        if len(plaintiff) == 0:
            plaintiff = ''
            return plaintiff
        else:
            for i in range(0, len(plaintiff)):
                plaintiff[i] = re.sub('[,，住所地]+', '', plaintiff[i])
                plaintiff[i] = re.sub('被告[0-9]{0,1}[：:]', '', plaintiff[i])
                plaintiff[i] = re.sub('[\W]', '', plaintiff[i])
            plaintiff = '，'.join(plaintiff)
            return plaintiff
    except:
        logger.info('---------------------------被告正则抽取失败')


# 原告
def find_plaintiff(text_str):
    """
    function: 原告信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的原告信息
    """
    text_str = text_str.replace(' ', '')
    try:
        plaintiff = re.findall(r'原告[0-9]{0,1}[：:].*?[,，住所地]+', text_str)
        if len(plaintiff) == 0:
            plaintiff = ''
            return plaintiff
        else:
            for i in range(0, len(plaintiff)):
                plaintiff[i] = re.sub('[,，住所地]+', '', plaintiff[i])
                plaintiff[i] = re.sub('原告[0-9]{0,1}[：:]', '', plaintiff[i])
                plaintiff[i] = re.sub('[\W]', '', plaintiff[i])
            plaintiff = '，'.join(plaintiff)
            return plaintiff
    except:
        logger.info('---------------------------原告正则抽取失败')


# 被上诉人
def find_appellee(text_str):
    """
    function: 被上诉人信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的被上诉人信息
    """
    text_str = text_str.replace(' ', '')
    try:
        plaintiff = re.findall(r'被上诉人[0-9]{0,1}[：:].*?[,，住所地]+', text_str)
        if len(plaintiff) == 0:
            plaintiff = ''
            return plaintiff
        else:
            for i in range(0, len(plaintiff)):
                plaintiff[i] = re.sub('[,，住所地]+', '', plaintiff[i])
                plaintiff[i] = re.sub('被上诉人[0-9]{0,1}[：:]', '', plaintiff[i])
                plaintiff[i] = re.sub('[\W]', '', plaintiff[i])
            plaintiff = '，'.join(plaintiff)
            return plaintiff
    except:
        logger.info('---------------------------被上诉人正则抽取失败')


# 上诉人
def find_appellor(text_str):
    """
    function: 上诉人信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的上诉人信息
    """
    text_str = text_str.replace(' ', '')
    try:
        plaintiff = re.findall(r'上诉人[0-9]{0,1}[：:].*?[,，住所地]+', text_str)
        if len(plaintiff) == 0:
            plaintiff = ''
            return plaintiff
        else:
            for i in range(0, len(plaintiff)):
                plaintiff[i] = re.sub('[,，住所地]+', '', plaintiff[i])
                plaintiff[i] = re.sub('上诉人[0-9]{0,1}[：:]', '', plaintiff[i])
                plaintiff[i] = re.sub('[\W]', '', plaintiff[i])
            plaintiff = '，'.join(plaintiff)
            return plaintiff
    except:
        logger.info('---------------------------上诉人正则抽取失败')


# 案号
def find_casecode(text_str):
    """
    function: 案号信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的案号信息
    """
    text_str = text_str.replace(' ', '')
    try:
        casecode = re.search(r'[\（\()]{1}[0-9]{4}[\）\)].*?[0-9]+[号]', text_str)
        if casecode == None:
            input_casecode = ''
            return input_casecode
        else:
            return casecode.group()
    except:
        logger.info('---------------------------案号正则抽取失败')


# 诉讼请求
def find_claim(text_str):
    """
    function: 诉讼请求信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的案号信息
    """
    text_str = text_str.replace(' ', '')
    try:
        result = re.search(r'原告[\u4e00-\u9fa5]+诉称[,].*?被告[\u4e00-\u9fa5]+辩称[，,]', text_str)
        if result == None:
            result2 = re.search(r'[上诉讼]{2}请求[：:1、]+.*?[。]事实[和与]{1}理由', text_str)
            if result2 == None:
                claim = ''
                return claim
            else:
                claim = re.sub('事实[\u4e00-\u9fa5]理由', '', result2.group())
                return claim
        else:
            claim = re.sub('被告[\u4e00-\u9fa5]+辩称[,]', '', result.group())
            return claim
    except:
        logger.info('---------------------------诉讼请求正则抽取失败')


# 判决结果
def find_judgment(text_str):
    """
    function: 判决结果信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的判决结果信息
    """
    text_str = text_str.replace(' ', '')
    try:
        result = re.search(r'判决如下[：:].*?审判[长员]{1}', text_str)
        if result == None:
            input_judgment = ''
            return input_judgment
        else:
            judgment = re.sub('审判[长员]{1}', '', result.group())
            return judgment
    except:
        logger.info('---------------------------判决结果正则抽取失败')


# 法院名称
def find_court1(text_str):
    """
    function: 法院名称信息抽取函数1
    param: 读入的文书文本
    return: 返回值为该文本抽取的法院名称信息
    """
    text_str = text_str.replace(' ', '')
    try:
        court = re.search(r'[\u4e00-\u9fa5]+人民法院', text_str)
        if court == None:
            input_court = ''
            return input_court
        else:
            return court.group()
    except:
        logger.info('---------------------------法院名称信息正则抽取失败')


def find_court2(text_str):
    """
    function: 法院名称信息抽取函数2
    param: 读入的文书文本
    return: 返回值为该文本抽取的法院名称信息
    """
    text_str = text_str.replace(' ', '')
    try:
        court = re.search(r'此致[\u4e00-\u9fa5]+人民法院', text_str)
        if court == None:
            input_court = ''
            return input_court
        else:
            court = re.sub('此致', '', court.group())
            court = re.sub('[\W]', '', court)
            return court
    except:
        logger.info('---------------------------法院名称信息正则抽取失败')


# 案件名称
def find_casename1(text_str):
    """
    function: 案件名称信息抽取函数1
    param: 读入的文书文本
    return: 返回值为该文本抽取的案件名称信息
    """
    text_str = text_str.replace(' ', '')
    try:
        casename = re.search(r'[:：][\u4e00-\u9fa5]+一案[,，]', text_str)
        if casename == None:
            input_casename = ''
            return input_casename
        else:
            casename = re.sub('[:,：，]', '', casename.group())
            return casename
    except:
        logger.info('---------------------------案件名称正则抽取失败')


def find_casename2(text_str):
    """
    function: 案件名称信息抽取函数2
    param: 读入的文书文本
    return: 返回值为该文本抽取的案件名称信息
    """
    text_str = text_str.replace(' ', '')
    try:
        casename = re.search(r'甲方因与[、\u4e00-\u9fa5]+一案', text_str)
        if casename == None:
            input_casename = ''
            return input_casename
        else:
            casename = re.sub('甲方因', '', casename.group())
            return casename
    except:
        logger.info('---------------------------案件名称正则抽取失败')


# 抬头银行名称
def find_bankname(text_str):
    """
    function: 抬头银行名称信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的抬头银行名称信息
    """
    text_str = text_str.replace(' ', '')
    try:
        bankname = re.search(r'[0-9]+[号][\u4e00-\u9fa5]+[:：]', text_str)
        if bankname == None:
            input_bankname = ''
            return input_bankname
        else:
            bankname1 = re.sub('[0-9]+号', '', bankname.group())
            bankname2 = re.sub('[:：]', '', bankname1)
            bankname2 = re.sub('[\W]', '', bankname1)
            return bankname2
    except:
        logger.info('---------------------------抬头银行名称正则抽取失败')


# 第三人
def find_third(text_str):
    """
    function: 第三人信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的第三人信息
    """
    text_str = text_str.replace(' ', '')
    try:
        third = re.findall(r'第三人[：:][\u4e00-\u9fa5]+[，,]', text_str)
        if len(third) == 0:
            third = ''
            return third
        else:
            for i in range(0, len(third)):
                third[i] = re.sub('[，,]', '', third[i])
                third[i] = re.sub('第三人[:：]', '', third[i])
                third[i] = re.sub('[\W]', '', third[i])
            third = '，'.join(third)
            return third
    except:
        logger.info('---------------------------第三人正则抽取失败')


# 事实与理由
def find_truth(text_str):
    """
    function: 事实与理由信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的事实与理由信息
    """
    text_str = text_str.replace(' ', '')
    try:
        result = re.search(r'事实与理由.*?此致', text_str)
        if result == None:
            input_truth = ''
            return input_truth
        else:
            truth = re.sub('事实与理由', '', result.group())
            truth = re.sub('此致', '', truth)
            return truth
    except:
        logger.info('---------------------------事实与理由正则抽取失败')


# 甲方
def find_partyA(text_str):
    """
    function: 甲方信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的甲方信息
    """
    text_str = text_str.replace(' ', '')
    try:
        partyA = re.findall(r'甲方[：:].*?[住所地|地址]', text_str)
        if len(partyA) == 0:
            partyA = ''
            return partyA
        else:
            for i in range(0, len(partyA)):
                partyA[i] = partyA[i][:len(partyA[i]) - 1]
                partyA[i] = re.sub('甲方[:：]', '', partyA[i])
                partyA[i] = re.sub('[\W]', '', partyA[i])
            partyA = '，'.join(partyA)
            return partyA
    except:
        logger.info('---------------------------甲方正则抽取失败')


# 乙方
def find_partyB(text_str):
    """
    function: 乙方信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的乙方信息
    """
    text_str = text_str.replace(' ', '')
    try:
        partyB = re.findall(r'乙方[：:].*?[住所地|地址]', text_str)
        if len(partyB) == 0:
            partyB = ''
            return partyB
        else:
            for i in range(0, len(partyB)):
                partyB[i] = partyB[i][:len(partyB[i]) - 1]
                partyB[i] = re.sub('乙方[:：]', '', partyB[i])
                partyB[i] = re.sub('[\W]', '', partyB[i])
            partyB = '，'.join(partyB)
            return partyB
    except:
        logger.info('---------------------------乙方正则抽取失败')


# 代理律师
def find_lawyers(text_str):
    """
    function: 代理律师信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的代理律师信息
    """
    text_str = text_str.replace(' ', '')
    try:
        lawyers = re.search(r'[委|指]派[、\u4e00-\u9fa5]+律师', text_str)
        if lawyers == None:
            input_lawyers = ''
            return input_lawyers
        else:
            lawyers = re.sub('[委|指]派', '', lawyers.group())
            lawyers = re.sub('律师', '', lawyers)
            lawyers = re.sub('[\W]', '', lawyers)
            return lawyers
    except:
        logger.info('---------------------------代理律师正则抽取失败')


# 代理权限
def find_authority(text_str):
    """
    function: 代理权限信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的代理权限信息
    """
    text_str = text_str.replace(' ', '')
    try:
        authority = re.search(r'委托代理权限.*?第三条', text_str)
        if authority == None:
            input_authority = ''
            return input_authority
        else:
            authority = re.sub('委托代理权限', '', authority.group())
            authority = re.sub('第三条', '', authority)
            return authority
    except:
        logger.info('---------------------------代理权限正则抽取失败')


# 费用
def find_cost(text_str):
    """
    function: 费用信息抽取函数
    param: 读入的文书文本
    return: 返回值为该文本抽取的费用信息
    """
    text_str = text_str.replace(' ', '')
    try:
        cost = re.search(r'律师代理费[0-9]+元[。]', text_str)
        if cost == None:
            input_cost = ''
            return input_cost
        else:
            cost = re.sub('律师代理费', '', cost.group())
            cost = re.sub('[。]', '', cost)
            cost = re.sub('[\W]', '', cost)
            return cost
    except:
        logger.info('---------------------------费用正则抽取失败')


# 日期
def find_date():
    """
    function: 日期信息抽取函数
    param: 读入的文书文本
    return: 暂时返回''，因为手写类日期暂时无法解决
    """
    input_date = ''
    return input_date


"""""""""""""""""""""""""""

主函数部分

"""""""""""""""""""""""""""


def main_call(filename):
    """
    function: 主功能入口
    param: 接收接口上传文件
    return: 返回最终的文字抽取结果、表格抽取结果
    """
    logger.info('---------------------------抽取功能开始运行---------------------------')
    root_path = os.getcwd()  # 获取当前根目录
    logger.info('--------------当前根目录：{}'.format(root_path))
    this_path = root_path + '/upload/'  # 合同文件夹
    doc_path = root_path + '/doctrans/'  # 用来转存doc文件的

    file_name = this_path + filename
    text_path = file_name  # 文档绝对位置
    table_result, end_result = [], [] # 结果集

    if file_name.endswith('pdf'):  # pdf格式
        logger.info('----------------pdf开始读取：{}'.format(filename))
        # 首先判断是否为图片型pdf，若是则调用ocr接口
        try:
            with pdfplumber.open(file_name) as pdf:
                # 读取页面，判断带有图片的页面比例是否超过50%
                img_count = 0
                pages_count = len(pdf.pages)
                for idx in range(pages_count):
                    if (is_img_page(pdf.pages[idx])):
                        img_count += 1
                if img_count / pages_count > 0.5:
                    # 获取ocr接口返回文本数据，进行抽取，并传出结果
                    ocr_text = ocr_interface(filename, file_name)
                    # 将读取到的pdf文字进行抽取
                    table_result, end_result = read_ocr_pdf(ocr_text, filename)
                    return table_result, end_result  # 返回结果列表
                else:
                    # 否则直接调用文字型pdf抽取
                    table_result, end_result = result_process_pdf(file_name, filename)  # 清晰结果并且存入csv文件
                    return table_result, end_result  # 返回结果列表
        except:
            logger.info('-----------------该文件抽取失败：{}'.format(file_name))

    elif file_name.endswith('docx'):  # docx格式
        try:
            table_result, end_result = result_process_docx(file_name, filename)  # 清晰结果并且存入csv文件
        except:
            logger.info('----------------该文件抽取失败：{}'.format(file_name))
        return table_result, end_result  # 返回结果列表

    elif file_name.endswith('doc'):  # doc格式
        try:
            doc_docx(text_path, doc_path, filename)
            table_result, end_result = result_process_docx(doc_path + filename + 'x', filename)  # 清晰结果并且存入csv文件
            return table_result, end_result  # 返回结果列表
        except:
            logger.info('-------------该文件抽取失败：{}'.format(file_name))

    logger.info('---------------------------本次抽取功能运行结束---------------------------')


if __name__ == '__main__':
    import time
    start_time = time.time()
    filename = 'cc_1.pdf'
    main_call(filename)
    end_time = time.time()
    print("Results (after {:.3f} seconds):".format(end_time - start_time))
